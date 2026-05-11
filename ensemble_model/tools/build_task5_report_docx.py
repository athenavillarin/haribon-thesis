from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

import pandas as pd


@dataclass(frozen=True)
class Paths:
    repo_root: Path
    ensemble_dir: Path
    results_dir: Path
    out_docx: Path


def _set_mono(run) -> None:
    run.font.name = "Consolas"


def _add_codeblock(doc, lines: list[str]) -> None:
    p = doc.add_paragraph()
    for i, ln in enumerate(lines):
        if i:
            p.add_run("\n")
        r = p.add_run(ln)
        _set_mono(r)


def _add_kv_paragraph(doc, key: str, value: str) -> None:
    p = doc.add_paragraph()
    p.add_run(f"{key}: ").bold = True
    p.add_run(value)


def _add_table(doc, df: pd.DataFrame, style: str = "Table Grid") -> None:
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = style
    hdr = table.rows[0].cells
    for j, col in enumerate(df.columns):
        hdr[j].text = str(col)
        for r in hdr[j].paragraphs[0].runs:
            r.bold = True
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(df.columns):
            val = row[col]
            cells[j].text = "" if pd.isna(val) else str(val)


def build_report(paths: Paths) -> None:
    from docx import Document

    per_split = pd.read_csv(paths.results_dir / "ensemble_per_split_metrics.csv")
    summary = pd.read_csv(paths.results_dir / "ensemble_summary.csv")
    obj2 = pd.read_csv(paths.results_dir / "obj2_model_comparison_final.csv")

    # Helper: split meta (use first row per split)
    split_meta = (
        per_split.sort_values(["split_num", "source", "name"])
        .drop_duplicates(subset=["split_num"])[
            ["split_num", "train_end", "test_start", "test_end", "n_train", "n_test", "positive_rate"]
        ]
        .sort_values("split_num")
        .reset_index(drop=True)
    )

    doc = Document()
    doc.core_properties.title = "Objective 2, Task 5 — Ensemble Model (Updated)"

    # Title block
    doc.add_paragraph("HARIBON Project").runs[0].bold = True
    doc.add_paragraph("Objective 2, Task 5 — Ensemble Model for HAB Detection").runs[0].bold = True
    doc.add_paragraph("Final Report (Updated / Fixed Ensemble Runner)")

    # 1. Overview
    doc.add_heading("1. Overview", level=1)
    doc.add_paragraph(
        "Task 5 closes Objective 2 of the HARIBON study by unifying the four independently trained HAB detection "
        "models into a single ensemble pipeline. The base models are XGBoost, LSTM, GRU, and a Transformer encoder."
    )
    doc.add_paragraph("Task 5 addresses three coordinated objectives:")
    for t in [
        "Combine XGBoost, LSTM, GRU, and Transformer via three ensemble strategies (soft vote, weighted average, stacking).",
        "Coordinate evaluation across all base models and ensemble variants on the shared 6-split rolling-origin framework.",
        "Compile the definitive Objective 2 results table (obj2_model_comparison_final.csv).",
    ]:
        doc.add_paragraph(t, style="List Bullet")

    doc.add_paragraph("The ensemble pipeline can be reproduced by running:")
    _add_codeblock(doc, ["cd ensemble_model", "python run_ensemble.py --transformer-scenario hybrid_adaptive"])

    # 2. Dataset and Feature Configuration
    doc.add_heading("2. Dataset and Feature Configuration", level=1)
    doc.add_heading("2.1 Source Dataset", level=2)
    _add_kv_paragraph(doc, "File", "final_compiled_dataset/Combined_Labeled.csv")
    doc.add_paragraph("Temporal coverage: 2016–2026 (daily observations). Target variable: red_tide_label (binary threshold ≥ 0.5).")

    doc.add_heading("2.2 Features (11)", level=2)
    feat_rows = [
        ("CHL", "Chlorophyll-a concentration", "mg/m³"),
        ("NDVI_daily", "Normalized Difference Vegetation Index", "dimensionless"),
        ("mlotst", "Mixed layer depth", "m"),
        ("precip_mm_day", "Daily precipitation", "mm/day"),
        ("so", "Sea water salinity", "PSU"),
        ("thetao", "Sea water potential temperature", "°C"),
        ("uo", "Eastward ocean current velocity", "m/s"),
        ("vo", "Northward ocean current velocity", "m/s"),
        ("wind_speed_ms", "Wind speed magnitude", "m/s"),
        ("wind_u_ms", "Eastward wind velocity component", "m/s"),
        ("wind_v_ms", "Northward wind velocity component", "m/s"),
    ]
    _add_table(doc, pd.DataFrame(feat_rows, columns=["Variable", "Description", "Unit"]))
    doc.add_paragraph("Table 1. Feature variables shared across all four base models.")

    doc.add_heading("2.3 Imputation Pipeline (Hybrid Gap‑Adaptive)", level=2)
    doc.add_paragraph(
        "Missing values are filled using the 4-phase Hybrid Gap‑Adaptive strategy (Objective 1’s selected method). "
        "Phases are applied in order until no missing values remain."
    )
    impute_rows = [
        ("1", "Linear interpolation", "Gaps ≤ 14 consecutive days per variable/location"),
        ("2", "Climatological mean", "Remaining gaps — Location × Month mean filled"),
        ("3", "Location mean fallback", "If climatological mean unavailable"),
        ("4", "Global mean (emergency)", "Final fallback to avoid NaNs"),
    ]
    _add_table(doc, pd.DataFrame(impute_rows, columns=["Phase", "Method", "Applied When"]))
    doc.add_paragraph("Table 2. 4-phase Hybrid Gap‑Adaptive imputation applied before split generation.")

    doc.add_heading("2.4 Rolling-Origin Split Framework (6 Splits)", level=2)
    _add_table(
        doc,
        split_meta.rename(
            columns=dict(
                split_num="Split",
                train_end="Train End",
                test_start="Test Start",
                test_end="Test End",
                n_train="Train Size",
                n_test="Test Size",
                positive_rate="Positive Rate",
            )
        ),
    )
    doc.add_paragraph("Table 3. Common 6-split rolling-origin evaluation framework.")

    # 3. Individual Model Baselines (from ensemble per-split outputs)
    doc.add_heading("3. Individual Model Baselines (6-Split Evaluation)", level=1)
    doc.add_paragraph(
        "The tables below report per-split performance from the ensemble runner on the aligned 6-split framework. "
        "AUC-ROC is primary; F1/precision/recall reflect the default threshold used during evaluation."
    )

    def model_table(name: str, title: str) -> None:
        doc.add_heading(title, level=2)
        sub = per_split[(per_split["source"] == "model") & (per_split["name"] == name)].copy()
        sub = sub.sort_values("split_num")
        out = sub[
            ["split_num", "train_end", "test_start", "test_end", "accuracy", "precision", "recall", "f1", "auc"]
        ].rename(
            columns=dict(
                split_num="Split",
                train_end="Train End",
                test_start="Test Start",
                test_end="Test End",
                accuracy="Accuracy",
                precision="Precision",
                recall="Recall",
                f1="F1",
                auc="AUC-ROC",
            )
        )
        # Compact formatting
        for c in ["Accuracy", "Precision", "Recall", "F1", "AUC-ROC"]:
            out[c] = out[c].map(lambda x: "" if pd.isna(x) else f"{float(x):.4f}")
        _add_table(doc, out)

    model_table("lstm", "3.1 LSTM — Rolling-Origin Results (6 Splits)")
    model_table("gru", "3.2 GRU — Rolling-Origin Results (6 Splits)")
    model_table("transformer", "3.3 Transformer — Rolling-Origin Results (6 Splits, hybrid_adaptive)")
    model_table("xgboost", "3.4 XGBoost — Rolling-Origin Results (6 Splits, refit-per-split)")

    # 4. Ensemble Results (summary)
    doc.add_heading("4. Ensemble Results (6 Splits, hybrid_adaptive)", level=1)
    doc.add_paragraph("Mean ± std across splits (ranked by AUC) from ensemble_model/results/ensemble_summary.csv.")
    summ = summary.copy()
    for c in [
        "accuracy_mean",
        "accuracy_std",
        "precision_mean",
        "precision_std",
        "recall_mean",
        "recall_std",
        "f1_mean",
        "f1_std",
        "auc_mean",
        "auc_std",
    ]:
        summ[c] = summ[c].map(lambda x: "" if pd.isna(x) else f"{float(x):.4f}")
    _add_table(
        doc,
        summ[
            [
                "rank_auc",
                "source",
                "name",
                "accuracy_mean",
                "precision_mean",
                "recall_mean",
                "f1_mean",
                "auc_mean",
                "n_splits",
            ]
        ].rename(
            columns=dict(
                rank_auc="Rank",
                source="Source",
                name="Model / Strategy",
                accuracy_mean="Accuracy (mean)",
                precision_mean="Precision (mean)",
                recall_mean="Recall (mean)",
                f1_mean="F1 (mean)",
                auc_mean="AUC (mean)",
                n_splits="n_splits",
            )
        ),
    )

    best = summary.sort_values("rank_auc").iloc[0]
    doc.add_paragraph(
        f"Best ensemble strategy by AUC: {best['name']} (AUC mean = {float(best['auc_mean']):.4f}, n_splits = {int(best['n_splits'])})."
    )

    # 5. Objective 2 Final Model Comparison
    doc.add_heading("5. Objective 2 Final Model Comparison", level=1)
    doc.add_paragraph("Generated from ensemble_model/results/obj2_model_comparison_final.csv.")
    _add_table(doc, obj2)

    # 6. Fixes applied (new)
    doc.add_heading("6. Fixes Applied in the Updated Ensemble Runner", level=1)
    for b in [
        "Transformer: can load split weights even if transformer_model/code/ modules are not present (notebook-only Transformer projects).",
        "LSTM: Keras 3 deserialization compatibility (drops quantization_config during Dense layer reconstruction).",
        "Windows: avoids console UnicodeEncodeError by using ASCII-only best-model print line.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    # 7. Outputs
    doc.add_heading("7. Output Files", level=1)
    _add_kv_paragraph(doc, "Per-split metrics", "ensemble_model/results/ensemble_per_split_metrics.csv")
    _add_kv_paragraph(doc, "Summary table", "ensemble_model/results/ensemble_summary.csv")
    _add_kv_paragraph(doc, "Final Objective 2 table", "ensemble_model/results/obj2_model_comparison_final.csv")

    paths.out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(paths.out_docx))


def main() -> None:
    repo_root = Path(__file__).resolve().parents[2]
    ensemble_dir = repo_root / "ensemble_model"
    results_dir = ensemble_dir / "results"
    out_docx = Path(r"c:\Users\meagie\Downloads\HARIBON_Task5_Report.docx")

    build_report(
        Paths(
            repo_root=repo_root,
            ensemble_dir=ensemble_dir,
            results_dir=results_dir,
            out_docx=out_docx,
        )
    )
    print(f"Wrote: {out_docx}")


if __name__ == "__main__":
    main()

