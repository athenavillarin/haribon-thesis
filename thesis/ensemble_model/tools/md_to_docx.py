from __future__ import annotations

import re
from pathlib import Path


def _set_mono(run) -> None:
    run.font.name = "Consolas"


def md_to_docx(md_path: Path, docx_path: Path) -> None:
    from docx import Document

    md_text = md_path.read_text(encoding="utf-8")
    lines = md_text.splitlines()

    doc = Document()
    doc.core_properties.title = md_path.stem

    in_code = False
    code_lines: list[str] = []

    def flush_code() -> None:
        nonlocal code_lines
        if not code_lines:
            return
        p = doc.add_paragraph()
        for i, ln in enumerate(code_lines):
            if i:
                p.add_run("\n")
            r = p.add_run(ln)
            _set_mono(r)
        code_lines = []

    bullet_re = re.compile(r"^\s*-\s+(.*)$")

    for raw in lines:
        line = raw.rstrip("\n")

        if line.strip().startswith("```"):
            if in_code:
                in_code = False
                flush_code()
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if not line.strip():
            doc.add_paragraph("")
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            continue

        m = bullet_re.match(line)
        if m:
            doc.add_paragraph(m.group(1).strip(), style="List Bullet")
            continue

        # Plain paragraph (keep backticks literal).
        doc.add_paragraph(line)

    if in_code:
        flush_code()

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))


if __name__ == "__main__":
    here = Path(__file__).resolve()
    root = here.parents[1]
    md = root / "ENSEMBLE_MODEL_REPORT.md"
    out = root / "ENSEMBLE_MODEL_REPORT.docx"
    md_to_docx(md, out)
    print(f"Wrote: {out}")

